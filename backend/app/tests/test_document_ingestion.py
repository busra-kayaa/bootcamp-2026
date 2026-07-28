from pathlib import Path

import pytest
from docx import Document

from app.pipelines.document_ingestion_pipeline import DocumentIngestionPipeline


@pytest.mark.asyncio
async def test_pipeline_reads_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "ornek.docx"
    document = Document()
    document.add_paragraph("Birinci paragraf")
    document.add_paragraph("İkinci paragraf")
    document.save(file_path)

    result = await DocumentIngestionPipeline().process(
        document_id=12, file_path=str(file_path)
    )

    assert result["document_id"] == 12
    assert result["page_count"] == 1
    assert result["extraction_method"] == "python-docx"
    assert result["character_count"] > 0
    assert result["chunks"]


@pytest.mark.asyncio
async def test_pipeline_rejects_unsupported_extension(tmp_path: Path) -> None:
    file_path = tmp_path / "ornek.txt"
    file_path.write_text("metin", encoding="utf-8")

    with pytest.raises(ValueError, match="Yalnızca PDF ve DOCX"):
        await DocumentIngestionPipeline().process(1, str(file_path))


@pytest.mark.asyncio
async def test_pipeline_rejects_zero_byte_pdf(tmp_path: Path) -> None:
    file_path = tmp_path / "bos.pdf"
    file_path.touch()

    with pytest.raises(ValueError, match="Dosya boş"):
        await DocumentIngestionPipeline().process(1, str(file_path))


@pytest.mark.asyncio
async def test_pipeline_rejects_fake_pdf(tmp_path: Path) -> None:
    file_path = tmp_path / "sahte.pdf"
    file_path.write_text("Bu bir PDF değil", encoding="utf-8")

    with pytest.raises(ValueError, match="geçerli bir PDF"):
        await DocumentIngestionPipeline().process(1, str(file_path))


@pytest.mark.asyncio
async def test_pipeline_rejects_docx_without_text(tmp_path: Path) -> None:
    file_path = tmp_path / "bos.docx"
    Document().save(file_path)

    with pytest.raises(ValueError, match="DOCX içerisinde okunabilir metin"):
        await DocumentIngestionPipeline().process(1, str(file_path))
