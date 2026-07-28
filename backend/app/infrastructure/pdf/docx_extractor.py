from pathlib import Path

from docx import Document

from .base import ExtractedPage


class DocxExtractor:
    """DOCX belgelerinden paragraf ve tablo metinlerini çıkarır."""

    def extract(self, file_path: str | Path) -> list[ExtractedPage]:
        path = Path(file_path)
        self._validate_file(path)

        try:
            document = Document(path)
        except Exception as exc:
            raise RuntimeError(f"DOCX dosyası okunurken hata oluştu: {path.name}") from exc

        parts = [paragraph.text.strip() for paragraph in document.paragraphs]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    parts.append(row_text)

        text = "\n".join(part for part in parts if part).strip()
        if not text:
            raise ValueError("DOCX içerisinde okunabilir metin bulunamadı.")

        return [
            {
                "page_number": 1,
                "text": text,
                "extraction_method": "python-docx",
            }
        ]

    @staticmethod
    def _validate_file(path: Path) -> None:
        if not path.exists():
            raise FileNotFoundError(f"DOCX dosyası bulunamadı: {path}")
        if not path.is_file():
            raise ValueError(f"Verilen yol bir dosyaya ait değil: {path}")
        if path.suffix.lower() != ".docx":
            raise ValueError("Desteklenmeyen dosya türü. Yalnızca DOCX kabul edilir.")
        if path.stat().st_size == 0:
            raise ValueError("DOCX dosyası boş.")
